"""
Utility functions for generating DOCX documents with formatted content.
"""
import os
import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# Helper to sanitize a file path into a base name without extension
def sanitize_filename(path: str) -> str:
    """
    Dado um caminho de arquivo (Windows ou Unix), retorna apenas o nome-base sem a extensão.
    Ex: "D:\\pasta\\LO.md"  -> "LO"
        "/home/user/ARQ.txt" -> "ARQ"
    """
    base = os.path.basename(path)
    name, _ = os.path.splitext(base)
    return name


def create_document(search_term, results):
    """
    Create a formatted DOCX document with search results.
    
    Args:
        search_term (str): The search term used for the query
        results (dict): Dictionary containing search results { source_path: [texts...] }
        
    Returns:
        Document: A python-docx Document object
    """
    # Create a new Document with custom styles
    doc = Document()
    
    # Set document properties
    core_props = doc.core_properties
    core_props.title = f'Search Results for: {search_term}'
    core_props.subject = 'Documento gerado a partir de busca lexical'
    
    # Set document margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Add title with custom style
    title = doc.add_heading(level=0)
    title_run = title.add_run(f'Search Results for: {search_term}')
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(16)
    title_run.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add a line break after title
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    
    # Add content for each source
    for file_path, matches in results.items():
        clean_name = sanitize_filename(file_path)
        
        # Add file heading
        file_heading = doc.add_heading(level=1)
        file_run = file_heading.add_run(clean_name)
        file_run.bold = True
        file_run.font.size = Pt(14)
        file_run.font.color.rgb = RGBColor(44, 62, 80)  # Dark blue-gray
        
        # Space after heading
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        
        # Add each match as numbered paragraph
        for i, match in enumerate(matches, start=1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.left_indent = Inches(0.5)
            
            # Number in bold
            run = p.add_run(f"{i}. ")
            run.bold = True
            
            # Add the actual text (with any markdown/formatted fragments)
            _add_formatted_text(p, match)
    
    # Add footer with page numbers
    _add_footer(doc)
    
    return doc




def _add_formatted_text(paragraph, text):
    """
    Adiciona texto formatado ao parágrafo, processando marcações Markdown.
    
    Args:
        paragraph: O parágrafo do documento
        text: Texto com formatação Markdown
    """
    # Define o alinhamento justificado para o parágrafo
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Processa blocos de código
    text = re.sub(r'```.*?\n(.*?)\n```', r'\1', text, flags=re.DOTALL)
    
    # Processa links
    text = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1 (\2)', text)
    
    # Processa listas
    text = re.sub(r'^\s*[-*+]\s+', '', text, flags=re.MULTILINE)
    
    # Processa blocos de citação
    text = re.sub(r'^>\s*', '', text, flags=re.MULTILINE)
    
    # Processa texto em negrito e itálico
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Texto em negrito
            run = paragraph.add_run(part[2:-2] + ' ')
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            # Texto em itálico
            run = paragraph.add_run(part[1:-1] + ' ')
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            # Texto em código
            run = paragraph.add_run(part[1:-1] + ' ')
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        elif part.strip():
            # Texto normal
            paragraph.add_run(part + ' ')





def _add_footer(doc):
    """
    Adiciona um rodapé com numeração de páginas ao documento.
    
    Args:
        doc: O documento ao qual adicionar o rodapé
    """
    section = doc.sections[0]
    footer = section.footer
    
    # Remove parágrafos padrão do rodapé
    for p in footer.paragraphs:
        p.clear()
    
    # Adiciona um parágrafo centralizado
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Adiciona o número da página
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    # Adiciona barra e número total de páginas
    run = paragraph.add_run(' / ')
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'NUMPAGES'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    # Estiliza o rodapé
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)



#--- Funções ---

#--- Funções ---


# 2. Converte arquivo Markdown em DOCX com Pandoc
#=============================================================================================
def convert_markdown_to_docx(markdown_text: str, docx_filename: str) -> bool:
    """
    Convert markdown text to a DOCX file using Pandoc.
    
    Args:
        markdown_text: String containing markdown content
        docx_filename: Path where the output DOCX file should be saved
        
    Returns:
        bool: True if conversion was successful, False otherwise
    """
    try:
        # Validate input
        if not markdown_text or not isinstance(markdown_text, str):
            print("Error: markdown_text must be a non-empty string")
            return False
            
        if not docx_filename or not isinstance(docx_filename, str):
            print("Error: docx_filename must be a non-empty string")
            return False
            
        # Ensure output directory exists
        output_dir = os.path.dirname(docx_filename)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # Run Pandoc command
        command = [
            "pandoc",
            "--from=markdown",
            "--to=docx",
            f"--output={docx_filename}",
            "--standalone"
        ]
        
        process = subprocess.Popen(
            command,
            stdin=subprocess.PIPE,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )
        
        # Send markdown content to Pandoc
        stdout, stderr = process.communicate(input=markdown_text.encode('utf-8'))
        
        if process.returncode != 0:
            print(f"Error converting markdown to DOCX: {stderr.decode('utf-8')}")
            return False
            
        if not os.path.exists(docx_filename):
            print("Error: Output file was not created")
            return False
            
        print(f"Successfully created: {os.path.abspath(docx_filename)}")
        return True
        
    except Exception as e:
        print(f"Unexpected error during markdown to DOCX conversion: {str(e)}")
        return False



#3. Aplica Macro ao DOCX gerado
#=============================================================================================
def aplicar_macro(docx_filename):
    # Converte para caminho absoluto para garantir que o Word encontre o arquivo
    abs_docx_path = os.path.abspath(docx_filename)
    
    # Nome exato da macro no Normal.dotm
    nome_macro = "Normal.Cons_IA.RAG_Clean"

    # Inicializa a biblioteca COM
    import pythoncom
    pythoncom.CoInitialize()
    
    try:
        # Inicia o Word
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False  # Torne True para ver a janela

        # Abre o documento com caminho absoluto
        print(f"Tentando abrir o arquivo: {abs_docx_path}")
        doc = word.Documents.Open(abs_docx_path)

        # Executa a macro do Normal.dotm
        try:
            word.Application.Run(nome_macro)
            print(f"Macro '{nome_macro}' executada com sucesso.")
        except Exception as e:
            print(f"Erro ao executar a macro: {e}")

        # Salva (como .docx)
        doc.SaveAs(abs_docx_path, FileFormat=12)  # 12 = .docx

        # Fecha tudo
        doc.Close(SaveChanges=True)
        return True
    except Exception as e:
        print(f"Erro ao processar o documento Word: {e}")
        return False
    finally:
        # Garante que o Word seja fechado mesmo em caso de erro
        try:
            word.Quit()
        except:
            pass
