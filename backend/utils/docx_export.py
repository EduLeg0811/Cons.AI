from io import BytesIO
import json
import logging
import unicodedata

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from html2docx import html2docx
import markdown2

logger = logging.getLogger(__name__)


# =======================
# CONSTANTES GERAIS
# =======================
DEFAULT_FONT_NAME   = "Calibri"
DEFAULT_FONT_SIZE   = Pt(10)             # corpo do texto
DEFAULT_FONT_COLOR  = RGBColor(0, 0, 0)  # preto

LINE_SPACING        = 1.0                # espaçamento entre linhas (simples)
SPACE_BEFORE        = Pt(0)              # espaço antes do parágrafo
SPACE_AFTER         = Pt(0)              # espaço depois do parágrafo

MARGIN_CM           = 2.5                # todas as margens em cm


def _add_paragraph_border_double(paragraph, color="000000", size_eights_pt=10):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'double')
        el.set(qn('w:sz'), str(size_eights_pt))
        el.set(qn('w:color'), color)
        el.set(qn('w:space'), '4')
        pBdr.append(el)
    pPr.append(pBdr)


def _add_bottom_border(paragraph, color="444444", size_eights_pt=10):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size_eights_pt))
    bottom.set(qn('w:color'), color)
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)







logger = logging.getLogger(__name__)








#_________________________________________________________  
# build_docx_deepdive
#_________________________________________________________  
def build_docx(data, group_results_by_book):

    # Extração de dados
    # =============================================================
    search_term   = data.get("search_term")
    search_type   = data.get("search_type")
    display_option = data.get("display_option")
    lexical       = data.get("lexical", [])
    semantical    = data.get("semantical", [])
    definologia   = data.get("definologia", "")
    descritivo    = data.get("descritivo", "")
    source_array  = data.get("source_array", [])
    max_results   = data.get("max_results", 10)

    # Garante string em definologia e descritivo
    defText = ""
    descrText = ""

    if isinstance(definologia, dict):
        defText = definologia.get("text") 
        if not isinstance(defText, str):
            defText = str(defText)

    if isinstance(descritivo, dict):
        descrText = descritivo.get("text") 
        if not isinstance(descrText, str):
            descrText = str(descrText)


    if search_type == "semantical":
        searchTypeTxt = "Semântica"
    elif search_type == "lexical" or search_type == "lexverb":
        searchTypeTxt = "Léxica"
    elif search_type == "deepdive":
        searchTypeTxt = "Deep Research"



    # Criação do documento
    # =============================================================
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = DEFAULT_FONT_NAME
    style.font.size = DEFAULT_FONT_SIZE
    style.font.color.rgb = DEFAULT_FONT_COLOR

    pf = style.paragraph_format
    pf.line_spacing  = LINE_SPACING
    pf.space_before  = SPACE_BEFORE
    pf.space_after   = SPACE_AFTER

    for section in doc.sections:
        section.top_margin    = Cm(MARGIN_CM)
        section.bottom_margin = Cm(MARGIN_CM)
        section.left_margin   = Cm(MARGIN_CM)
        section.right_margin  = Cm(MARGIN_CM)



    # 1. Cabeçalho
    # =============================================================
    p = doc.add_paragraph(str(search_term).capitalize())
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(20)
    run.font.bold = True
    _add_paragraph_border_double(p, size_eights_pt=9)

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Tipo de pesquisa
    p = doc.add_paragraph()
    p.add_run("Tipo de pesquisa: ").bold = True
    p.add_run(searchTypeTxt)
    doc.add_paragraph("")

    # Termo
    p = doc.add_paragraph()
    p.add_run("Termo de pesquisa: ").bold = True
    p.add_run(str(search_term).capitalize())
    doc.add_paragraph("")

    # Source array
    p = doc.add_paragraph()
    p.add_run("Livros pesquisados: ").bold = True
    p.add_run(str(source_array))
    doc.add_paragraph("")

    # Display Option
    p = doc.add_paragraph()
    p.add_run("Display Option: ").bold = True
    p.add_run(display_option)
    doc.add_paragraph("")

  
    # 2. Estatística geral
    # =============================================================
    stats = {}

    logger.info("data.items(): " + str(data.items()))

    for key, value in data.items():
        if isinstance(value, list) and key in ("lexical", "semantical") and len(value) > 0:
            stats[key] = len(value)
            logger.info("stats[key]: " + str(stats[key]))

    total_count = sum(stats.values())
    p = doc.add_paragraph()
    p.add_run("Estatística de Resultados: ").bold = True
    p.add_run(str(total_count))
    for src, items in stats.items():
        if items > 0:
            doc.add_paragraph(f"• {src}: {items}")

            logger.info("Estatística de Resultados: " + str(items))

    doc.add_paragraph("")
    doc.add_paragraph("")
   
    logger.info("Estatística de Resultados (total): " + str(total_count))




    # 3. Definologia e Descritivo
    # =============================================================

    if defText or descrText:
        # badge_def = doc.add_paragraph()
        # run = badge_def.add_run("Definologia")
        # run.font.size = Pt(12)
        # run.font.bold = True
        # run.font.color.rgb = RGBColor(0, 0, 0)
        # badge_def.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # divider_def = doc.add_paragraph()
        # _add_bottom_border(divider_def)
        # doc.add_paragraph("")

        if defText:
            p = doc.add_paragraph()
            p.add_run("Definição: ").bold = True
            insert_markdown_into_paragraph(str(defText).capitalize(), p)
            doc.add_paragraph("")
        
        if descrText:
            p = doc.add_paragraph()
            p.add_run("Descritivos: ").bold = True
            insert_markdown_into_paragraph(str(descrText).capitalize(), p)
            doc.add_paragraph("")


    doc.add_paragraph("")        


    # 4 Display Option
    # =============================================================
    
    # display_mode == "simple" : semantical data grouped by source or sorted by score (NORMAL)
    # ----------------------------------------------------------------------------------------
    if (display_option == "simple"):
        display_results_simple(doc, data, search_type, group_results_by_book)

    # display_mode == "advanced" : semantical data grouped by source or sorted by score (NORMAL)
    # ----------------------------------------------------------------------------------------
    elif (display_option == "unified"):
        display_results_unified(doc, data, search_type)



    # 6. Referências
    # =============================================================
    doc.add_paragraph("")

    # linha separadora antes da legenda
    divider_ref = doc.add_paragraph()
    _add_bottom_border(divider_ref)
    

    # título da legenda
    p = doc.add_paragraph()
    p.add_run("\nLegenda:\n").bold = True

    if search_type == "semantical" or search_type == "deepdive":
        p.add_run("@ :  ").bold = True 
        p.add_run("score de similaridade semântica (0 = máxima correspondência)\n")
    
    p.add_run("# :  ").bold = True 
    p.add_run("nº do parágrafo (Livros); nº do verbete (EC); nº da questão (Conscienciograma)")

    doc.add_paragraph("")


    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()



#_________________________________________________________
# display_results_unified       
#_________________________________________________________
def display_results_unified(doc, data, search_type):

    # 1. Extrai arrays já ordenados separadamente
    array_data_lexical = flatten_data(
        data, fields=["lexical"], sort_fields=[("number", "crescent")]
    )
    array_data_semantical = flatten_data(
        data, fields=["semantical"], sort_fields=[("score", "decrescent")]
    )

    newData = {}

    def _normalize_source(value):
        if isinstance(value, list):
            for candidate in value:
                if candidate:
                    return str(candidate)
            return "Unknown"
        if value in (None, ""):
            return "Unknown"
        return str(value)

    def _add_items(items, field):
        for original in items or []:
            if not isinstance(original, dict):
                continue
            src_value = original.get("source") or original.get("source_array")
            src = _normalize_source(src_value)
            obj = dict(original)
            obj.setdefault("source", src)
            obj.setdefault("field", field)
            newData.setdefault(src, {"lexical": [], "semantical": []})[field].append(obj)

    # separa lexical e semantical por fonte
    _add_items(array_data_lexical, "lexical")
    _add_items(array_data_semantical, "semantical")

    # eliminar duplicados pelo campo "text"
    def _normalize_text(txt):
        if not txt:
            return ""
        return " ".join(txt.lower().split())

    for src, groups in newData.items():
        for field in ["lexical", "semantical"]:
            seen = set()
            unique_items = []
            for it in groups[field]:
                txt = (
                    it.get("text")
                    or it.get("content_text")
                    or it.get("markdown")
                    or (it.get("metadata") or {}).get("markdown")
                    or ""
                )
                norm = _normalize_text(txt)
                if norm not in seen:
                    seen.add(norm)
                    unique_items.append(it)
            newData[src][field] = unique_items

    # mostrar resultados por fonte
    for src, groups in newData.items():
        if not groups["lexical"] and not groups["semantical"]:
            continue

        # Subtítulo da fonte
        doc.add_paragraph("")
        src_title = doc.add_paragraph()
        run = src_title.add_run(bookName(src))
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 100, 200)
        src_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph("")
        doc.add_paragraph("")

        counter = 0

        # Primeiro lexical
        for it in groups["lexical"]:
            counter += 1
            mdText = (
                (it.get("metadata") or {}).get("markdown")
                or it.get("markdown")
                or it.get("content_text")
                or it.get("text")
                or ""
            )
            p = doc.add_paragraph()
            num_run = p.add_run(f"{counter}. ")
            num_run.bold = True
            num_run.font.color.rgb = RGBColor(0, 0, 200)
            num_run.font.size = Pt(9)

            if src in ("EC", "ECALL_DEF", "ECWV", "ECALL"):
                mdText = f"**Definologia.** {mdText}"

            insert_markdown_into_paragraph(mdText, p)

            metaInfo = createMetaInfo(it, src)
            metaInfo_p = doc.add_paragraph(metaInfo)
            metaInfo_p.runs[0].font.size = Pt(8)
            metaInfo_p.runs[0].font.color.rgb = RGBColor(150, 0, 50)

            doc.add_paragraph("")

        # Linha divisória entre lexical e semantical
        if groups["lexical"] and groups["semantical"]:
            divider_p = doc.add_paragraph()
            _add_bottom_border(divider_p, color="888888", size_eights_pt=8)
            doc.add_paragraph("")
        else:
            doc.add_paragraph("")

        # Depois semantical
        for it in groups["semantical"]:
            counter += 1
            mdText = (
                (it.get("metadata") or {}).get("markdown")
                or it.get("markdown")
                or it.get("content_text")
                or it.get("text")
                or ""
            )
            p = doc.add_paragraph()
            num_run = p.add_run(f"{counter}. ")
            num_run.bold = True
            num_run.font.color.rgb = RGBColor(0, 0, 200)
            num_run.font.size = Pt(9)

            if src in ("EC", "ECALL_DEF", "ECWV", "ECALL"):
                mdText = f"**Definologia.** {mdText}"
            if src == "LO" and search_type == "semantical":
                mdText = f"**{it.get('title')}**. {mdText}"

            insert_markdown_into_paragraph(mdText, p)

            metaInfo = createMetaInfo(it, src)
            metaInfo_p = doc.add_paragraph(metaInfo)
            metaInfo_p.runs[0].font.size = Pt(8)
            metaInfo_p.runs[0].font.color.rgb = RGBColor(150, 0, 50)

            doc.add_paragraph("")









#_________________________________________________________
# display_results_simple       
#_________________________________________________________
def display_results_simple(doc, data, search_type, group_results_by_book):


    if search_type == "lexical":
        array_search_loop = ["lexical"]
    elif search_type == "lexverb":
        array_search_loop = ["lexical"]
    elif search_type == "semantical":
        array_search_loop = ["semantical"]
    elif search_type == "deepdive":
        array_search_loop = ["lexical", "semantical"]


    # Itera nos modos de pesquisa (lexical, semantical) ou ambos (deepdive)
    for current_search_mode in array_search_loop:

        doc.add_paragraph("")
        doc.add_paragraph("")


    # ---------------------------------------------------------
    # 1. Lexical
    # ---------------------------------------------------------
    if current_search_mode == "lexical":

        search_mode_txt = "Pesquisa Léxica"
        array_data = flatten_data(
            data, fields=["lexical"], sort_fields=[("number", "crescent")]
        )


        if array_data:

            
            badge_p = doc.add_paragraph()
            run = badge_p.add_run("Resultados da Pesquisa Léxica")
            
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 100, 0)
            badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            divider_p = doc.add_paragraph()
            _add_bottom_border(divider_p, size_eights_pt=9)
            doc.add_paragraph("")
            doc.add_paragraph("")



            if (group_results_by_book):

                counter = 1
                
                # Agrupar resultados por source
                # -----------------------------
                grouped = {}
                for it in array_data:
                    src = it.get("source") or it.get("source_array") or "Unknown"
                    grouped.setdefault(src, []).append(it)

                # Iterar por fonte
                for src, items in grouped.items():

                    doc.add_paragraph("")
                    doc.add_paragraph("")
        
                    # Subtítulo com o nome da fonte
                    src_title = doc.add_paragraph()
                    run = src_title.add_run(bookName(src))
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 100, 200)
                    src_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    doc.add_paragraph("")
                    doc.add_paragraph("")

                    # Linha divisória
                    # divider_src = doc.add_paragraph()
                    # _add_bottom_border(divider_src)
                    # doc.add_paragraph("")

                    # Iterar pelos itens daquela fonte
                    for it in items:
                        
                        mdText = (
                            (it.get("metadata") or {}).get("markdown")
                            or it.get("markdown")
                            or it.get("content_text")
                            or it.get("text")
                            or ""
                        )


                        # Novo parágrafo com número
                        p = doc.add_paragraph()
                        num_run = p.add_run(f"{counter}. ")
                        num_run.bold = True
                        num_run.font.color.rgb = RGBColor(0, 0, 200)
                        num_run.font.size = Pt(9)

                        # Ajustes especiais (livros, verbetes, etc)
                        if src in ('EC', 'ECALL_DEF', 'ECWV', 'ECALL'):
                            mdText = f"**Definologia.** {mdText}"

                        if src == 'LO' and search_type == 'semantical':
                            mdText = f"**{it.get('title')}**. {mdText}"

                        insert_markdown_into_paragraph(mdText, p)

                        metaInfo = createMetaInfo(it, src)
                        metaInfo_p = doc.add_paragraph(metaInfo)
                        metaInfo_p.runs[0].font.size = Pt(8)
                        metaInfo_p.runs[0].font.color.rgb = RGBColor(150, 0, 50)
                        metaInfo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        doc.add_paragraph("")
                        counter += 1


            # ELSE: sem agrupamento
            # ---------------------
            else:

                doc.add_paragraph("")
                doc.add_paragraph("")
                

                for it in array_data:
                    src    = it.get("source") or it.get("source_array") or ""
                    mdText = (
                        (it.get("metadata") or {}).get("markdown")
                        or it.get("markdown")
                        or it.get("content_text")
                        or it.get("text")
                        or ""
                    )

                    # Novo parágrafo com número
                    p = doc.add_paragraph()
                    num_run = p.add_run(f"{counter}. ")
                    num_run.bold = True
                    num_run.font.color.rgb = RGBColor(0, 0, 200)
                    num_run.font.size = Pt(9)

                    # Se for verbete da EC, antepor "Definologia."
                    if src in ('EC', 'ECALL_DEF', 'ECWV', 'ECALL'):
                        mdText = f"**Definologia.** {mdText}"

                    # Insere o texto Markdown neste parágrafo
                    insert_markdown_into_paragraph(mdText, p)


                    # 5. Metadados
                    # =============================================================
                    metaInfo = createMetaInfo(it, src)
                    metaInfo_p = doc.add_paragraph(metaInfo)
                    metaInfo_p.runs[0].font.size = Pt(8)
                    metaInfo_p.runs[0].font.color.rgb = RGBColor(150, 0, 50)
                    metaInfo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    doc.add_paragraph("")
                    counter += 1



    # ---------------------------------------------------------
    # 2. Semantic
    # ---------------------------------------------------------
    if current_search_mode == "semantical":

        search_mode_txt = "Pesquisa Semântica"
        array_data = flatten_data(
            data, fields=["semantical"], sort_fields=[("score", "crescent")]
        )


        if array_data:

            badge_p = doc.add_paragraph()
            run = badge_p.add_run("Resultados da " + search_mode_txt+ " (ordenados por similaridade)")
    
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 100, 0)
            badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            #divider_p = doc.add_paragraph()
            divider_p = doc.add_paragraph()
            _add_bottom_border(divider_p, size_eights_pt=10)

            doc.add_paragraph("")
            doc.add_paragraph("")
            

            # Agrupa por livros (simple)
            # ==========================
            if (group_results_by_book):
                
            
                grouped = {}
                for it in array_data:
                    src = it.get("source") or it.get("source_array") or "Unknown"
                    grouped.setdefault(src, []).append(it)

                # Iterar por fonte
                for src, items in grouped.items():

                    doc.add_paragraph("")
                    doc.add_paragraph("")
        
                    # Subtítulo com o nome da fonte
                    src_title = doc.add_paragraph()
                    run = src_title.add_run(bookName(src))
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 100, 200)
                    src_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    doc.add_paragraph("")
                    doc.add_paragraph("")

                    # Linha divisória
                    # divider_src = doc.add_paragraph()
                    # _add_bottom_border(divider_src)
                    # doc.add_paragraph("")

                    # Iterar pelos itens daquela fonte
                    for it in items:
                        
                        mdText = (
                            (it.get("metadata") or {}).get("markdown")
                            or it.get("markdown")
                            or it.get("content_text")
                            or it.get("text")
                            or ""
                        )

                        # Novo parágrafo com número
                        p = doc.add_paragraph()
                        num_run = p.add_run(f"{counter}. ")
                        num_run.bold = True
                        num_run.font.color.rgb = RGBColor(0, 0, 200)
                        num_run.font.size = Pt(9)

                        # Ajustes especiais
                        if src in ('EC', 'ECALL_DEF', 'ECWV', 'ECALL'):
                            mdText = f"**Definologia.** {mdText}"

                        if src == 'LO' and search_type == 'semantical':
                            mdText = f"**{it.get('title')}**. {mdText}"

                        insert_markdown_into_paragraph(mdText, p)

                        metaInfo = createMetaInfo(it, src)
                        metaInfo_p = doc.add_paragraph(metaInfo)
                        metaInfo_p.runs[0].font.size = Pt(8)
                        metaInfo_p.runs[0].font.color.rgb = RGBColor(150, 0, 50)
                        metaInfo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        doc.add_paragraph("")
                        counter += 1


            # ELSE: sem agrupamento
            # ---------------------
            else:

                doc.add_paragraph("")
                doc.add_paragraph("")
                

                for it in array_data:
                    src    = it.get("source") or it.get("source_array") or ""
                    mdText = (
                        (it.get("metadata") or {}).get("markdown")
                        or it.get("markdown")
                        or it.get("content_text")
                        or it.get("text")
                        or ""
                    )

                    # Novo parágrafo com número
                    p = doc.add_paragraph()
                    num_run = p.add_run(f"{counter}. ")
                    num_run.bold = True
                    num_run.font.color.rgb = RGBColor(0, 0, 200)
                    num_run.font.size = Pt(9)


                    # Se for verbete da EC, antepor "Definologia."
                    if src in ('EC', 'ECALL_DEF', 'ECWV', 'ECALL'):
                        mdText = f"**Definologia.** {mdText}"

                    # Se for pensata do LO em Semantical, antepor Titulo em negrito antes do texto
                    if src in ('LO'):
                        mdText = f"**{it.get('title')}**. {mdText}"


                    # Insere o texto Markdown neste parágrafo
                    insert_markdown_into_paragraph(mdText, p)


                    # 5. Metadados
                    # =============================================================
                    metaInfo = createMetaInfo(it, src)
                    metaInfo_p = doc.add_paragraph(metaInfo)
                    metaInfo_p.runs[0].font.size = Pt(8)
                    metaInfo_p.runs[0].font.color.rgb = RGBColor(150, 0, 50)
                    metaInfo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    doc.add_paragraph("")
                    counter += 1                                                                                    











#_________________________________________________________
# insert_markdown_into_paragraph
#_________________________________________________________
def insert_markdown_into_paragraph(mdText, paragraph):
    """
    Converte Markdown para runs e insere no parágrafo existente (paragraph).
    Mantém negrito, itálico e sublinhado.
    """
    if not mdText:
        return

    # Converte Markdown -> HTML
    html = markdown2.markdown(mdText)

    # Cria doc temporário a partir do HTML
    tmp_bytes = html2docx(html, title="tmp")
    tmp_doc = Document(tmp_bytes)

    # Copia runs do doc temporário para o parágrafo existente
    for para in tmp_doc.paragraphs:
        for run in para.runs:
            new_run = paragraph.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.size = run.font.size or Pt(10)
            new_run.font.color.rgb = (
                run.font.color.rgb if run.font.color else RGBColor(0, 0, 0)
            )
        
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.space_before  = 0
        paragraph.space_after   = 0
        paragraph.line_spacing = LINE_SPACING
    









#_________________________________________________________  
# process_deepdive
#_________________________________________________________  
def flatten_data(
    data,
    fields=["lexical", "semantical"],
    sort_fields=[("number", "crescent")],
    ignore_accents=False
):

    if isinstance(data, str):
        try:
            data = json.loads(data)
        except Exception as e:
            logger.error(f"process_deepdive_teste: data inválido ({e})")
            return []

    lista = []
    for field in fields:
        if field in data:
            if isinstance(data[field], list):
                for item in data[field]:
                    if isinstance(item, dict):
                        obj = dict(item)
                        obj["field"] = field
                        if obj.get("metadata") is None:
                            obj["metadata"] = {}
                        lista.append(obj)
                    else:
                        lista.append({"value": item, "field": field})
            else:
                logger.debug(f"process_deepdive_teste --- {field} não é lista ({type(data[field])})")

    def normalize_string(val: str) -> str:
        s = str(val).lower()
        if ignore_accents:
            s = "".join(c for c in unicodedata.normalize("NFD", s)
                        if unicodedata.category(c) != "Mn")
        return s

    def convert_value(val):
        if val is None:
            return None
        try:
            return float(val)
        except (ValueError, TypeError):
            return normalize_string(val)

    for field, order in reversed(sort_fields):
        reverse = order.lower() in ["decrescent", "desc", "descending"]

        def key_func(obj, f=field):
            val = convert_value(obj.get(f, None))
            if val is None:
                return float("inf") if not reverse else float("-inf")
            return val

        lista.sort(key=key_func, reverse=reverse)

    return lista




#______________________________________________________________________________________________
# createMetaInfo
#______________________________________________________________________________________________
def createMetaInfo(it, src):
    meta = it.get("metadata", {}) or {}

    mt_book   = bookName(src) + " | "
    mt_number = it.get("number") or meta.get("number") or ""
    mt_title  = it.get("title") or meta.get("title") or ""
    mt_score  = it.get("score") or meta.get("score") or ""
    mt_author = it.get("author") or meta.get("author") or ""
    mt_date   = it.get("date") or meta.get("date") or ""
    mt_section= it.get("section") or meta.get("section") or ""
    mt_theme  = it.get("theme") or meta.get("theme") or ""
    mt_area   = it.get("area") or meta.get("area") or ""
    mt_folha  = it.get("folha") or meta.get("folha") or ""

    # Adiciona " | " apenas se valor existir
    def fmt(val): return f"{val} | " if val else ""

    mt_number = fmt(mt_number)
    mt_title  = fmt(mt_title)
    mt_score  = fmt(mt_score)
    mt_author = fmt(mt_author)
    mt_date   = fmt(mt_date)
    mt_section= fmt(mt_section)
    mt_theme  = fmt(mt_theme)
    mt_area   = fmt(mt_area)
    mt_folha  = fmt(mt_folha)

    metaInfo_list = [mt_book]

    if src == "LO":
        metaInfo_list.append(mt_title)
    elif src == 'DAC':
        metaInfo_list.extend([mt_title, mt_author, mt_date, mt_section])
    elif src == 'CCG':
        metaInfo_list.extend([mt_title, mt_section, mt_folha])
    elif src in ('EC', 'ECALL_DEF', 'ECWV', 'ECALL'):
        metaInfo_list.extend([mt_title, mt_area, mt_theme, mt_author, mt_date])
    else:
        metaInfo_list.append(mt_title)

    # number e score sempre no final
    if mt_number:
        metaInfo_list.append("#" + mt_number)
    if mt_score:
        metaInfo_list.append("@" + mt_score)

    # Remove vazios e junta
    metaInfo = "".join([s for s in metaInfo_list if s])

    # Remove " | " final se sobrou
    if metaInfo.endswith(" | "):
        metaInfo = metaInfo[:-3]

    return metaInfo



      

#______________________________________________________________________________________________
# bookName  --- call from [bridge.js] <call_llm>
#______________________________________________________________________________________________
def bookName(source):

    realName = source;

    if (source == 'HSR'):
        realName = 'Homo sapiens reurbanisatus';

    if (source == 'HSP'):
        realName = 'Homo sapiens pacificus';
    
    if (source == '200TEAT'):
        realName = '200 Teáticas da Conscienciologia';

    if (source == '700EXP'):
        realName = '700 Experimentos da Conscienciologia';
    
    if (source == 'TEMAS'):
        realName = 'Temas da Conscienciologia';
    
    if (source == 'PROEXIS'):
        realName = 'Manual da Proéxis';
    
    if (source == 'TNP'):
        realName = 'Manual da Tenepes';
    
    if (source == 'DUPLA'):
        realName = 'Manual da Dupla Evolutiva';
    
    if (source == 'LO'):
        realName = 'Léxico de Ortopensatas';
    
    if (source == 'EC' or source == 'ECALL' or source == 'ECALL_DEF'):
        realName = 'Enciclopédia da Conscienciologia';
    
    if (source == 'DAC'):
        realName = 'Dicionário de Argumentos da Conscienciologia';
    
    if (source == 'PROJ'):
        realName = 'Projeciologia';

    if (source == 'CCG'):
        realName = 'Conscienciograma';

    return realName;






